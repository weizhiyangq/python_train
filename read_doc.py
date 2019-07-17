# -*- coding: utf-8 -*-
"""
Created on Mon Mar 11 13:13:19 2019

@author: YWZQ
"""

import docx


import pandas as pd

file_write=open('ganjv1.txt',mode='w+')
doc=docx.Document('ganjv1.docx')  #注意docx只能读取docx文档不能读取doc文档
print(len(doc.paragraphs))
try:
    for para in doc.paragraphs:
        print(para.text)
        file_write.write(para.text)
        file_write.write('\n')
except UnicodeEncodeError:
    pass
    